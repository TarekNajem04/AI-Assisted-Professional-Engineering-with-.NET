"""
SPDX-License-Identifier: MIT
Copyright (c) 2025 Tarek Najem

This file is part of the AI-Assisted Professional Engineering with .NET
book project (https://github.com/TarekNajem04/AI-Assisted-Professional-Engineering-with-.NET).
It is subject to the terms and conditions of the MIT License as published
in the LICENSE file at the root of this repository.

This header must not be removed or modified without preserving the
LICENSE file reference and copyright notice.

Description:
    Insert full-page cover images into a DOCX file.
    Front cover goes on the first page; back cover on the last page.
    Uses ECMA-376-compliant sectPr placement for each section.
"""
import argparse, os, sys
from docx import Document
from docx.shared import Emu, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'shared'))
import log

A4_W_EMU = int(Cm(21))
A4_H_EMU = int(Cm(29.7))
_TWIP_EMU = 635  # 1 twip = 635 EMU
_A4_W_TWIP = str(round(Cm(21) / _TWIP_EMU))      # 11906
_A4_H_TWIP = str(round(Cm(29.7) / _TWIP_EMU))    # 16838


def _img_size(path):
    with open(path, 'rb') as f:
        hdr = f.read(32)
    if hdr[6:10] in (b'JFIF', b'Exif'):
        return _jpg_size(path)
    if hdr.startswith(b'\x89PNG\r\n\x1a\n'):
        return int.from_bytes(hdr[16:20], 'big'), int.from_bytes(hdr[20:24], 'big')
    return (794, 1123)


def _jpg_size(path):
    import struct
    with open(path, 'rb') as f:
        f.seek(2)
        while True:
            b = f.read(4)
            if len(b) < 4: break
            mrk, blk = struct.unpack('>HH', b)
            if mrk in (0xFFC0, 0xFFC2):
                h, w = struct.unpack('>HH', f.read(5)[1:5])
                return w, h
            f.seek(blk - 2, 1)
    return (794, 1123)


def _disp(img_w, img_h):
    pa = A4_W_EMU / A4_H_EMU
    ia = img_w / img_h
    if abs(ia - pa) < 0.001:
        return A4_W_EMU, A4_H_EMU
    if ia > pa:
        return A4_W_EMU, int(A4_W_EMU / ia)
    return int(A4_H_EMU * ia), A4_H_EMU


def _normalise_sectPr(body):
    """
    Pandoc embeds <w:sectPr> inside the last <w:p>'s <w:pPr>.
    Extract it to be a direct <w:body> child so python-docx sections work.
    """
    last_child = list(body)[-1] if len(body) > 0 else None
    if last_child is not None:
        nested = last_child.find(qn('w:pPr') + '/' + qn('w:sectPr'))
        if nested is not None:
            pPr = nested.getparent()
            pPr.remove(nested)
            body.append(nested)
            log.log_info("  [Normalised] sectPr extracted from paragraph pPr to direct body child")


def _cover_para(doc, img_path):
    """Add a cover image paragraph, return it."""
    iw, ih = _img_size(img_path)
    dw, dh = _disp(iw, ih)
    para = doc.add_paragraph()
    para.alignment = 1
    pf = para.paragraph_format
    pf.space_before = Emu(0)
    pf.space_after = Emu(0)
    pf.line_spacing = 1.0
    run = para.add_run()
    run.add_picture(img_path, width=Emu(dw), height=Emu(dh))
    return para


def _zero_sectpr():
    """Create a <w:sectPr> with A4 page and zero margins."""
    sp = OxmlElement('w:sectPr')
    sz = OxmlElement('w:pgSz')
    sz.set(qn('w:w'), _A4_W_TWIP)
    sz.set(qn('w:h'), _A4_H_TWIP)
    sp.append(sz)
    pm = OxmlElement('w:pgMar')
    for a in ('top','right','bottom','left','header','footer','gutter'):
        pm.set(qn(f'w:{a}'), '0')
    sp.append(pm)
    return sp


def _embed_sectPr_in_last_para(body, sectPr):
    """
    Move a direct-body-child sectPr into the LAST paragraph before it.
    sectPr goes inside w:pPr so python-docx's section parser finds it.
    """
    orig_children = list(body)
    orig_idx = orig_children.index(sectPr)
    body.remove(sectPr)
    children = list(body)
    start = min(orig_idx - 1, len(children) - 1)
    for i in range(start, -1, -1):
        if children[i].tag == qn('w:p'):
            p = children[i]
            pPr = p.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p.insert(0, pPr)
            pPr.append(sectPr)
            return True
    body.insert(min(orig_idx, len(body)), sectPr)
    return False


def _embed_direct_sectPrs(body):
    """
    After adding covers, find ALL direct-body-child sectPrs that are
    NOT the last child of body, and embed them in the last preceding
    paragraph. This makes the DOCX ECMA-376 compliant.
    """
    children = list(body)
    direct_sectPrs = [c for c in children if c.tag == qn('w:sectPr')]
    if not direct_sectPrs:
        return
    last_sectPr = children[-1]
    if last_sectPr.tag != qn('w:sectPr'):
        return  # last child is not a sectPr — nothing to do
    for sp in direct_sectPrs:
        if sp is last_sectPr:
            continue  # keep the LAST sectPr as a direct child (correct for final section)
        if _embed_sectPr_in_last_para(body, sp):
            log.log_info("  [Embedded] section break sectPr into last paragraph")


def add_front_cover(doc, img_path):
    """
    Insert front cover as FIRST section.
    The zero-margin sectPr is EMBEDDED as the last child of the cover
    paragraph (ECMA-376 compliant for non-final sections).
    """
    if not os.path.isfile(img_path):
        log.log_warn(f"Front cover not found: {log.rel(img_path)}")
        return
    body = doc.element.body
    para = _cover_para(doc, img_path)
    sp = _zero_sectpr()
    para_elem = para._element

    # Put sectPr in w:pPr so python-docx's section parser finds it
    pPr = para_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para_elem.insert(0, pPr)
    pPr.append(sp)

    # Move paragraph to position 0
    body.remove(para_elem)
    body.insert(0, para_elem)

    log.log_info(f"[COVER] Front cover added: {log.rel(img_path)}")


def add_back_cover(doc, img_path):
    """
    Insert back cover as LAST section.
    The zero-margin sectPr is appended as the LAST child of <body>.
    This is correct per ECMA-376 for the final section.
    """
    if not os.path.isfile(img_path):
        log.log_warn(f"Back cover not found: {log.rel(img_path)}")
        return
    body = doc.element.body
    para = _cover_para(doc, img_path)
    sp = _zero_sectpr()
    para_elem = para._element

    body.remove(para_elem)
    body.append(para_elem)
    body.append(sp)

    log.log_info(f"[COVER] Back cover added: {log.rel(img_path)}")


def main():
    parser = argparse.ArgumentParser(description="Add cover images to DOCX")
    parser.add_argument("docx_path")
    parser.add_argument("--front")
    parser.add_argument("--back")
    parser.add_argument("--repo-root")
    parser.add_argument("--config")
    args = parser.parse_args()
    log.init(config_path=args.config, repo_root=args.repo_root or '')
    if not args.front and not args.back:
        log.log_info("Nothing to do")
        return
    if not os.path.isfile(args.docx_path):
        log.log_error(f"DOCX not found: {log.rel(args.docx_path)}")
        sys.exit(1)

    doc = Document(args.docx_path)
    body = doc.element.body
    _normalise_sectPr(body)

    if args.back:
        add_back_cover(doc, args.back)
    if args.front:
        add_front_cover(doc, args.front)

    # Make non-last section breaks ECMA-376 compliant (embedded in paragraphs)
    _embed_direct_sectPrs(body)

    doc.save(args.docx_path)

    # Verify margins on ALL sectPrs (direct + embedded)
    doc2 = Document(args.docx_path)
    body2 = doc2.element.body

    all_sectPrs = []
    all_sectPrs.extend(body2.findall(qn('w:sectPr')))
    all_sectPrs.extend(body2.findall('.//' + qn('w:p') + '/' + qn('w:sectPr')))
    all_sectPrs.extend(body2.findall('.//' + qn('w:pPr') + '/' + qn('w:sectPr')))

    for i, sp in enumerate(all_sectPrs):
        pm = sp.find(qn('w:pgMar'))
        if pm is not None:
            t = int(pm.get(qn('w:top'), '0')) / 567
            b = int(pm.get(qn('w:bottom'), '0')) / 567
            l = int(pm.get(qn('w:left'), '0')) / 567
            r = int(pm.get(qn('w:right'), '0')) / 567
            log.log_info(f"  Section {i}: T={t:.1f} B={b:.1f} L={l:.1f} R={r:.1f} cm")
    log.log_ok(f"Covers added to: {log.rel(args.docx_path)}")


if __name__ == "__main__":
    main()

