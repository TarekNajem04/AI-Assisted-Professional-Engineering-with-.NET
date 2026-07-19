"""
SPDX-License-Identifier: MIT
Copyright (c) 2025 Tarek Najem

This file is part of the AI-Assisted Professional Engineering with .NET
book project (https://github.com/TarekNajem04/AI-Assisted-Professional-Engineering-with-.NET).

Description:
    Generate a Pandoc reference DOCX template with the default
    style profile — Constantia serif body (12pt), Segoe UI headings (28/18/14),
    teal accent #005563, cool blue-gray code blocks.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
import json


def _set_font(style, name, size_pt=11, bold=False, italic=False, color=None):
    f = style.font
    f.name = name
    f.size = Pt(size_pt)
    f.bold = bold
    f.italic = italic
    if color:
        f.color.rgb = RGBColor(*color)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{name}" w:hAnsi="{name}" w:cs="{name}" w:eastAsia="{name}"/>')
        rpr.insert(0, rFonts)
    else:
        for a in ('ascii', 'hAnsi', 'cs', 'eastAsia'):
            rFonts.set(qn(f'w:{a}'), name)


def _set_pf(style, sb=0, sa=0, ls=1.15, align=None, fli=None, li=None, ri=None, keep=False):
    pf = style.paragraph_format
    pf.space_before = Pt(sb)
    pf.space_after = Pt(sa)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = ls
    if align:
        pf.alignment = align
    if fli is not None:
        pf.first_line_indent = Cm(fli)
    if li is not None:
        pf.left_indent = Cm(li)
    if ri is not None:
        pf.right_indent = Cm(ri)
    if keep:
        pf.keep_with_next = True


def _pandoc_style(doc, name, style_type=WD_STYLE_TYPE.PARAGRAPH):
    if name in [s.name for s in doc.styles]:
        return doc.styles[name]
    return doc.styles.add_style(name, style_type)


def _add_ph_shading(pPr, fill_hex):
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill_hex}"/>'))


def _add_ph_left_border(pPr, sz=8, color="C0C8D4"):
    b = f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="{sz}" w:space="6" w:color="{color}"/></w:pBdr>'
    pPr.append(parse_xml(b))


def _rel(path, root):
    if not root:
        return os.path.basename(path)
    p = path.replace('\\', '/')
    r = root.replace('\\', '/')
    if p.startswith(r):
        return p[len(r):].lstrip('/')
    return os.path.basename(path)


def create_docx_template(output_path, config=None, project_root=None):
    doc = Document()

    GD1 = (30, 30, 30)
    GM1 = (60, 60, 60)
    GL1 = (100, 100, 100)
    GL2 = (140, 140, 140)
    TD  = (0, 85, 105)
    TD2 = (0, 100, 130)
    AC  = (81, 43, 212)
    BODY = 'Constantia'
    HEAD = 'Segoe UI'

    sec = doc.sections[0]
    sec.page_width  = Cm(21.0)
    sec.page_height = Cm(29.7)
    if config and 'Docx' in config and 'Margins' in config['Docx']:
        m = config['Docx']['Margins']
        for s in doc.sections:
            s.top_margin    = Cm(m.get('Top', 2.0))
            s.bottom_margin = Cm(m.get('Bottom', 2.5))
            s.left_margin   = Cm(m.get('Left', 2.8))
            s.right_margin  = Cm(m.get('Right', 1.8))
    elif config and 'docx' in config and 'margins' in config['docx']:
        m = config['docx']['margins']
        for s in doc.sections:
            s.top_margin    = Cm(m.get('top', 2.0))
            s.bottom_margin = Cm(m.get('bottom', 2.5))
            s.left_margin   = Cm(m.get('left', 2.8))
            s.right_margin  = Cm(m.get('right', 1.8))
    else:
        for s in doc.sections:
            s.top_margin    = Cm(2.0)
            s.bottom_margin = Cm(2.5)
            s.left_margin   = Cm(2.8)
            s.right_margin  = Cm(1.8)
    sec.header_distance = Cm(1.2)
    sec.footer_distance = Cm(1.5)
    sec.different_first_page_header_footer = True

    # ── Modify Pandoc-default styles ──

    n = _pandoc_style(doc, 'Normal')
    _set_font(n, BODY, 12, color=GD1)
    _set_pf(n, sa=6, ls=1.15)

    bt = _pandoc_style(doc, 'Body Text')
    _set_font(bt, BODY, 12, color=GD1)
    _set_pf(bt, sa=6, ls=1.15, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    h1 = _pandoc_style(doc, 'Heading 1')
    _set_font(h1, HEAD, 28, bold=True, color=TD)
    _set_pf(h1, sb=48, sa=20, ls=1.1, keep=True)

    h2 = _pandoc_style(doc, 'Heading 2')
    _set_font(h2, HEAD, 18, bold=True, color=GD1)
    _set_pf(h2, sb=30, sa=14, ls=1.15, keep=True)

    h3 = _pandoc_style(doc, 'Heading 3')
    _set_font(h3, HEAD, 14, bold=True, color=GM1)
    _set_pf(h3, sb=22, sa=10, ls=1.15, keep=True)

    sc = _pandoc_style(doc, 'Source Code')
    _set_font(sc, 'Consolas', 9, color=GD1)
    _set_pf(sc, sb=14, sa=14, ls=1.0, li=0.4)
    pPr_sc = sc.element.get_or_add_pPr()
    pPr_sc.set(qn('w:bidi'), '0')
    _add_ph_shading(pPr_sc, "EBF0F5")
    _add_ph_left_border(pPr_sc, 10, "005563")

    ttl = _pandoc_style(doc, 'Title')
    _set_font(ttl, HEAD, 28, bold=True, color=TD)
    _set_pf(ttl, sb=48, sa=20, ls=1.1, keep=True)

    cap = _pandoc_style(doc, 'Caption')
    _set_font(cap, 'Segoe UI', 10, italic=True, color=GL1)
    _set_pf(cap, sa=10, ls=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)

    cf = _pandoc_style(doc, 'CaptionedFigure')
    _set_font(cf, BODY, 12, color=GD1)
    _set_pf(cf, align=WD_ALIGN_PARAGRAPH.CENTER)

    icap = _pandoc_style(doc, 'ImageCaption')
    _set_font(icap, HEAD, 10, bold=True, color=GL1)
    _set_pf(icap, sa=10, ls=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Override python-docx default blue styles ──
    for name in ['Subtitle', 'Heading 4', 'Heading 5', 'Heading 6', 'Heading 7', 'Heading 8',
                 'Heading 1 Char', 'Heading 2 Char', 'Heading 3 Char', 'Heading 4 Char',
                 'Heading 5 Char', 'Heading 6 Char', 'Heading 7 Char', 'Heading 8 Char',
                 'Subtitle Char', 'Intense Quote', 'Intense Quote Char', 'Intense Emphasis']:
        try:
            s = _pandoc_style(doc, name)
            _set_font(s, 'Segoe UI', 10, color=GL2)
        except:
            pass

    # ── Custom Book styles ──

    bb = _pandoc_style(doc, 'BookBody')
    _set_font(bb, BODY, 12, color=GD1)
    _set_pf(bb, sa=6, ls=1.15, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    bh1 = _pandoc_style(doc, 'BookHeading1')
    _set_font(bh1, HEAD, 28, bold=True, color=TD)
    _set_pf(bh1, sb=48, sa=20, ls=1.1, keep=True)

    bh2 = _pandoc_style(doc, 'BookHeading2')
    _set_font(bh2, HEAD, 18, bold=True, color=GD1)
    _set_pf(bh2, sb=30, sa=14, ls=1.15, keep=True)

    bh3 = _pandoc_style(doc, 'BookHeading3')
    _set_font(bh3, HEAD, 14, bold=True, color=GM1)
    _set_pf(bh3, sb=22, sa=10, ls=1.15, keep=True)

    bc = _pandoc_style(doc, 'BookCode')
    _set_font(bc, 'Consolas', 9, color=GD1)
    _set_pf(bc, sb=14, sa=14, ls=1.0, li=0.4)
    pPr_bc = bc.element.get_or_add_pPr()
    pPr_bc.set(qn('w:bidi'), '0')
    _add_ph_shading(pPr_bc, "EBF0F5")
    _add_ph_left_border(pPr_bc, 10, "005563")

    bq = _pandoc_style(doc, 'BookQuote')
    _set_font(bq, BODY, 12, italic=True, color=GM1)
    _set_pf(bq, sb=12, sa=12, ls=1.15, li=1.5, ri=1.5)
    _add_ph_left_border(bq.element.get_or_add_pPr(), 14, f"{TD[0]:02X}{TD[1]:02X}{TD[2]:02X}")
    _add_ph_shading(bq.element.get_or_add_pPr(), "F0F5F8")

    btab = _pandoc_style(doc, 'BookTable', WD_STYLE_TYPE.TABLE)
    _set_font(btab, BODY, 10, color=GD1)
    tbl_pr = btab.element.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        btab.element.append(tbl_pr)
    tbl_pr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="8" w:space="0" w:color="C0C8D4"/>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="C0C8D4"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E4EA"/>'
        f'</w:tblBorders>'
    ))

    bcap = _pandoc_style(doc, 'BookCaption')
    _set_font(bcap, HEAD, 10, italic=True, color=GL1)
    _set_pf(bcap, sa=10, ls=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)

    bic = _pandoc_style(doc, 'BookInlineCode', WD_STYLE_TYPE.CHARACTER)
    _set_font(bic, 'Consolas', 9.5, color=TD)
    bic.element.get_or_add_rPr().set(qn('w:rtl'), '0')

    bd = _pandoc_style(doc, 'BookDate')
    _set_font(bd, BODY, 11, italic=True, color=GL2)
    _set_pf(bd, sa=48, ls=1.2, align=WD_ALIGN_PARAGRAPH.CENTER)

    bicap = _pandoc_style(doc, 'BookImageCaption')
    _set_font(bicap, HEAD, 10, italic=True, color=GL1)
    _set_pf(bicap, sb=6, sa=12, ls=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)

    btoc1 = _pandoc_style(doc, 'BookTOC1')
    _set_font(btoc1, HEAD, 12, bold=True, color=GD1)
    _set_pf(btoc1, sb=6, sa=2, ls=1.3)

    btoc2 = _pandoc_style(doc, 'BookTOC2')
    _set_font(btoc2, HEAD, 11, color=GD1)
    _set_pf(btoc2, sb=2, sa=2, ls=1.2, li=1.5)

    btoc3 = _pandoc_style(doc, 'BookTOC3')
    _set_font(btoc3, HEAD, 10, color=GM1)
    _set_pf(btoc3, sb=1, sa=1, ls=1.2, li=3.0)

    # ── Header / Footer ──

    sec.first_page_header.is_linked_to_previous = False
    sec.header.is_linked_to_previous = False
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("AI-Assisted Professional Engineering with .NET")
    hr.font.name = HEAD
    hr.font.size = Pt(8)
    hr.font.color.rgb = RGBColor(130, 140, 150)

    sec.first_page_footer.is_linked_to_previous = False
    sec.footer.is_linked_to_previous = False
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    fp._p.append(parse_xml(
        f'<w:r {nsdecls("w")}><w:rPr><w:sz w:val="18"/><w:color w:val="646E78"/>'
        f'<w:rFonts w:ascii="Source Sans Pro"/></w:rPr>'
        f'<w:t xml:space="preserve">\u2014 </w:t></w:r>'
    ))
    fp._p.append(parse_xml(
        f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE ">'
        f'<w:r><w:rPr><w:sz w:val="18"/><w:color w:val="646E78"/>'
        f'<w:rFonts w:ascii="Source Sans Pro"/></w:rPr><w:t>1</w:t></w:r>'
        f'</w:fldSimple>'
    ))
    fp._p.append(parse_xml(
        f'<w:r {nsdecls("w")}><w:rPr><w:sz w:val="18"/><w:color w:val="646E78"/>'
        f'<w:rFonts w:ascii="Source Sans Pro"/></w:rPr>'
        f'<w:t xml:space="preserve"> \u2014</w:t></w:r>'
    ))

    # ── Sample content ──

    doc.add_paragraph("AI-Assisted Professional Engineering with .NET", style='Title')
    doc.add_paragraph("Tarek Najem", style='BookDate')
    doc.add_paragraph(
        "This technical reference explores the intersection of generative "
        "artificial intelligence and professional software engineering ...",
        style='Body Text')
    doc.add_paragraph(
        "A second paragraph with Constantia body at 12pt (1.15 line spacing), "
        "Segoe UI headings at 28/18/14pt in teal. Cool blue-gray code blocks "
        "with a teal left accent border. This is the Default profile.",
        style='Body Text')
    doc.add_page_break()
    doc.add_paragraph("Table of Contents", style='Heading 1')
    doc.add_paragraph("Chapter 1: Foundations", style='Heading 1')
    doc.add_paragraph("Section 1.1: The Engineering Mindset", style='Heading 2')
    doc.add_paragraph("Subsection: Critical Thinking", style='Heading 3')
    doc.add_paragraph(
        "Professional software engineering requires a disciplined approach "
        "to decision-making. Every architectural choice involves tradeoffs "
        "between competing concerns such as performance, maintainability, "
        "scalability, and security.",
        style='Body Text')
    doc.add_paragraph(
        'public static class Program\n{\n    public static void Main()\n    {\n'
        '        Console.WriteLine("Hello, .NET World!");\n'
        '    }\n}',
        style='Source Code')
    p = doc.add_paragraph("Inline code is ")
    p.add_run("Console.WriteLine", style='BookInlineCode')
    p.add_run(" inside body text.")

    doc.add_paragraph("Figure 1: System architecture diagram", style='BookImageCaption')
    doc.add_paragraph("Chapter 1: Foundations", style='BookTOC1')
    doc.add_paragraph("Section 1.1: Overview", style='BookTOC2')
    doc.add_paragraph("Historical Context", style='BookTOC3')

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    print(f"  [PROFILE] Reference DOCX created: {_rel(output_path, project_root)}")


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument('--output', default=None)
    parser.add_argument('--config', default=None)
    parser.add_argument('--project-root', default=None)
    args = parser.parse_args()
    config = None
    if args.config:
        with open(args.config, 'r', encoding='utf-8') as f:
            config = json.load(f)
    if args.output:
        output_path = args.output
    else:
        output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "custom-reference.docx")
    create_docx_template(output_path, config=config, project_root=args.project_root)
